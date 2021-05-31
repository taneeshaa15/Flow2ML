import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import pandas as pd
import io

class Process_Csv:
	'''
    Class containing methods to analyse non image datasets.
	'''
	def __init__(self,df):
		'''
		Initializes various attributes regarding to the object.
		Args : 
        	df : (Dataframe) Python dataframe to be analysed.
		'''
		self.df = df
		try:                        
			# try creating a results directory
			p = os.getcwd()
			p = os.path.join(p,"GeneratedReports")
			self.results_path = p
			if not os.path.exists(self.results_path):
				os.mkdir(self.results_path)
		except Exception as e:
			print(f"Unable to create directory for results due to {e}.")
		self.create_analysis_docx()
	
	def add_table_to_doc(self, table, doc):
		'''
		Helper function to format a table and add it to a document.
		Args : 
        	table : (Dataframe) Python dataframe to be added.
			doc: (Document) Document to add the table to.
		'''
		t = doc.add_table(table.shape[0] + 1, table.shape[1])
		t.style = 'Table Grid'
		# add the header rows.
		for j in range(table.shape[-1]):
			t.cell(0, j).text = table.columns[j]

		# add the rest of the table frame
		for i in range(table.shape[0]):
			for j in range(table.shape[-1]):
				t.cell(i + 1, j).text = str(table.values[i, j])

	def create_analysis_docx(self):
		'''
		Creates the analysis report.
		Args : 
        	None
		'''
		document =  Document()
		document.add_heading('EDA Report')

		# add df.head() 
		document.add_heading('df.head', level = 2)
		data = self.df.head()
		self.add_table_to_doc(data, document)

		# add df.describe()
		document.add_heading('df.describe', level = 2)
		data = self.df.describe()
		data.insert(0, 'Statistic', data.index)
		self.add_table_to_doc(data, document)

		# add df.info()
		buffer = io.StringIO()
		document.add_heading('df.info', level = 2)
		self.df.info(buf = buffer)
		data = buffer.getvalue()
		document.add_paragraph(data)

		# add df.nunique
		document.add_heading('df.nunique', level = 2)
		data = self.df.nunique(axis = 0).to_frame()
		data.insert(0, 'Column', data.index)
		data.columns = ['Column', 'Number of unique values']
		self.add_table_to_doc(data, document)

		# add missing values
		document.add_heading('df.isnull', level = 2)
		data = self.df.isnull().sum().to_frame()
		data.insert(0, 'Column', data.index)
		data.columns = ['Column', 'Number of missing values values']
		self.add_table_to_doc(data, document)
		document.save(os.path.join(self.results_path, 'csv_analysis.docx'))